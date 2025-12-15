import os
from typing import List
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT

from app.models import Profile, Grant
from app.services.firestore_service import FirestoreService


class PacketGeneratorService:
    """Service for generating grant application packets"""
    
    def __init__(self, firestore_service: FirestoreService):
        self.firestore_service = firestore_service
        self.output_dir = "packets"
        
        # Create output directory if it doesn't exist
        os.makedirs(self.output_dir, exist_ok=True)
    
    async def generate_packet(
        self,
        profile_id: str,
        grant_ids: List[str],
        format: str = "pdf"
    ) -> str:
        """
        Generate a grant application packet
        Returns: filepath to the generated packet
        """
        # Get profile
        profile = await self.firestore_service.get_profile(profile_id)
        if not profile:
            raise ValueError(f"Profile {profile_id} not found")
        
        # Get grants
        grants = []
        for grant_id in grant_ids:
            grant = await self.firestore_service.get_grant(grant_id)
            if grant:
                grants.append(grant)
        
        if not grants:
            raise ValueError("No valid grants found")
        
        # Generate packet based on format
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"grant_packet_{profile_id}_{timestamp}"
        
        if format.lower() == "pdf":
            filepath = os.path.join(self.output_dir, f"{filename}.pdf")
            self._generate_pdf(profile, grants, filepath)
        elif format.lower() == "docx":
            filepath = os.path.join(self.output_dir, f"{filename}.docx")
            self._generate_docx(profile, grants, filepath)
        else:
            raise ValueError(f"Unsupported format: {format}")
        
        return filepath
    
    def _generate_docx(self, profile: Profile, grants: List[Grant], filepath: str):
        """Generate DOCX packet"""
        doc = Document()
        
        # Title
        title = doc.add_heading("Grant Application Packet", 0)
        title.alignment = 1  # Center
        
        # Profile information
        doc.add_heading("Applicant Information", 1)
        doc.add_paragraph(f"Name: {profile.name}")
        doc.add_paragraph(f"Email: {profile.email}")
        doc.add_paragraph(f"Persona: {profile.persona}")
        doc.add_paragraph(f"Region: {profile.region}")
        if profile.gpa:
            doc.add_paragraph(f"GPA: {profile.gpa}")
        doc.add_paragraph(f"Income Level: {profile.income_level}")
        doc.add_paragraph("")
        
        # Grants section
        doc.add_heading("Recommended Grants", 1)
        
        for i, grant in enumerate(grants, 1):
            doc.add_heading(f"{i}. {grant.title}", 2)
            doc.add_paragraph(f"Organization: {grant.organization}")
            doc.add_paragraph(f"Amount: ${grant.amount:,.2f}")
            doc.add_paragraph(f"Deadline: {grant.deadline}")
            doc.add_paragraph(f"Description: {grant.description}")
            
            if grant.url:
                doc.add_paragraph(f"Application URL: {grant.url}")
            
            # Eligibility criteria
            doc.add_paragraph("Eligibility Criteria:")
            if grant.eligible_personas:
                doc.add_paragraph(f"  • Personas: {', '.join(grant.eligible_personas)}")
            if grant.eligible_regions:
                doc.add_paragraph(f"  • Regions: {', '.join(grant.eligible_regions)}")
            if grant.min_gpa:
                doc.add_paragraph(f"  • Minimum GPA: {grant.min_gpa}")
            if grant.income_requirements:
                doc.add_paragraph(f"  • Income Levels: {', '.join(grant.income_requirements)}")
            
            doc.add_paragraph("")
            
            if i < len(grants):
                doc.add_page_break()
        
        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        footer.alignment = 1  # Center
        
        doc.save(filepath)
    
    def _generate_pdf(self, profile: Profile, grants: List[Grant], filepath: str):
        """Generate PDF packet"""
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#333333',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor='#333333',
            spaceAfter=12,
            spaceBefore=12
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=14,
            textColor='#555555',
            spaceAfter=10,
            spaceBefore=10
        )
        
        # Title
        story.append(Paragraph("Grant Application Packet", title_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Profile information
        story.append(Paragraph("Applicant Information", heading_style))
        story.append(Paragraph(f"<b>Name:</b> {profile.name}", styles['Normal']))
        story.append(Paragraph(f"<b>Email:</b> {profile.email}", styles['Normal']))
        story.append(Paragraph(f"<b>Persona:</b> {profile.persona}", styles['Normal']))
        story.append(Paragraph(f"<b>Region:</b> {profile.region}", styles['Normal']))
        if profile.gpa:
            story.append(Paragraph(f"<b>GPA:</b> {profile.gpa}", styles['Normal']))
        story.append(Paragraph(f"<b>Income Level:</b> {profile.income_level}", styles['Normal']))
        story.append(Spacer(1, 0.3 * inch))
        
        # Grants section
        story.append(Paragraph("Recommended Grants", heading_style))
        story.append(Spacer(1, 0.2 * inch))
        
        for i, grant in enumerate(grants, 1):
            story.append(Paragraph(f"{i}. {grant.title}", subheading_style))
            story.append(Paragraph(f"<b>Organization:</b> {grant.organization}", styles['Normal']))
            story.append(Paragraph(f"<b>Amount:</b> ${grant.amount:,.2f}", styles['Normal']))
            story.append(Paragraph(f"<b>Deadline:</b> {grant.deadline}", styles['Normal']))
            story.append(Paragraph(f"<b>Description:</b> {grant.description}", styles['Normal']))
            
            if grant.url:
                story.append(Paragraph(f"<b>Application URL:</b> {grant.url}", styles['Normal']))
            
            # Eligibility criteria
            story.append(Paragraph("<b>Eligibility Criteria:</b>", styles['Normal']))
            if grant.eligible_personas:
                story.append(Paragraph(f"• Personas: {', '.join(grant.eligible_personas)}", styles['Normal']))
            if grant.eligible_regions:
                story.append(Paragraph(f"• Regions: {', '.join(grant.eligible_regions)}", styles['Normal']))
            if grant.min_gpa:
                story.append(Paragraph(f"• Minimum GPA: {grant.min_gpa}", styles['Normal']))
            if grant.income_requirements:
                story.append(Paragraph(f"• Income Levels: {', '.join(grant.income_requirements)}", styles['Normal']))
            
            if i < len(grants):
                story.append(PageBreak())
            else:
                story.append(Spacer(1, 0.3 * inch))
        
        # Footer
        footer_text = f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=10,
            textColor='#666666',
            alignment=TA_CENTER
        )
        story.append(Paragraph(footer_text, footer_style))
        
        doc.build(story)
